import os
import json
import mimetypes
import zipfile
import time
import random
import hashlib
from datetime import date
from io import BytesIO

import streamlit as st
from google import genai
from google.genai import types

# PDF renderer with no system deps
import fitz  # PyMuPDF

# for DOCX output
from docx import Document             # pip install python-docx
# for PDF output
from fpdf import FPDF                 # pip install fpdf

# ←── PAGE & THEME CONFIG ──────────────────────────────────────────────────────
st.set_page_config(
    page_title="Handwritten OCR",
    page_icon="🖋️",
    layout="centered",
    initial_sidebar_state="expanded"
)

st.markdown(
    """
    <style>
      :root {--primary-color:#004080;--bg-color:#f0f2f6;--sec-bg-color:#ffffff;--text-color:#333;}
      .main, .reportview-container {background-color:var(--bg-color)!important;color:var(--text-color)!important;}
      .block-container {padding:1.5rem 2rem!important;}
      h1 {font-size:3rem!important;color:var(--primary-color)!important;}
      h2 {font-size:2rem!important;color:var(--primary-color)!important;}
      .stMarkdown, .stText {font-size:18px!important;}
      .stTextArea>div>div>textarea {font-size:16px!important;}
      .stButton>button {
        background-color:var(--primary-color)!important;
        color:#fff!important;
        font-size:16px!important;
        padding:0.5rem 1rem!important;
      }
      .sidebar .sidebar-content {
        background-color:var(--sec-bg-color)!important;
        color:var(--text-color)!important;
        font-size:18px!important;
      }
    </style>
    """,
    unsafe_allow_html=True,
)

# ←── API & USAGE ───────────────────────────────────────────────────────────────
# DEFAULT_API_KEY = " "
DEFAULT_API_KEY = st.secrets["genai"]["api_key"]
USAGE_FILE = "usage.json"
DAILY_LIMIT = 5

def guess_mime(fname: str) -> str:
    mime, _ = mimetypes.guess_type(fname)
    return mime or "application/octet-stream"

# ---------------------- RESILIENT OCR (retries + fallbacks) -------------------
def _sha1(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()

_OCR_CACHE = {}

def ocr_with_gemini(client: genai.Client, data: bytes, fname: str, prompt: str) -> str:
    """
    Robust OCR call:
    - retries on 503/UNAVAILABLE with exponential backoff + jitter
    - falls back across several Gemini vision models
    - caches identical image bytes within the session
    """
    cache_key = (_sha1(data), prompt)
    if cache_key in _OCR_CACHE:
        return _OCR_CACHE[cache_key]

    models = [
        "gemini-2.5-flash",
        "gemini-1.5-flash",
        "gemini-1.5-flash-8b",
    ]

    last_err = None
    for model in models:
        base = 0.8  # base backoff seconds
        for attempt in range(6):  # ~0.8s → ~25s worst-case for a single model
            try:
                part = types.Part.from_bytes(data=data, mime_type=guess_mime(fname))
                resp = client.models.generate_content(
                    model=model,
                    contents=[part, prompt],
                )
                text = (getattr(resp, "text", "") or "").strip()
                _OCR_CACHE[cache_key] = text
                return text
            except Exception as e:
                msg = str(e)
                if ("503" in msg) or ("UNAVAILABLE" in msg) or ("deadline" in msg.lower()):
                    # transient → retry with backoff + jitter
                    sleep_s = base * (2 ** attempt) + random.uniform(0, 0.5)
                    time.sleep(sleep_s)
                    last_err = e
                    continue
                # non-transient → bubble up
                raise
        # exhausted retries for this model; try next
        last_err = last_err or RuntimeError(f"Model {model} failed after retries.")
    # all models failed
    raise last_err or RuntimeError("All OCR attempts failed.")

def load_usage():
    if os.path.isfile(USAGE_FILE):
        with open(USAGE_FILE, "r") as f:
            d = json.load(f)
    else:
        d = {"date": "", "count": 0}
    today = date.today().isoformat()
    if d.get("date") != today:
        d = {"date": today, "count": 0}
    return d

def save_usage(d):
    with open(USAGE_FILE, "w") as f:
        json.dump(d, f)

def make_pdf(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.splitlines():
        pdf.multi_cell(0, 10, line)
    return pdf.output(dest="S").encode("latin1")

def make_docx(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

st.title("Handwritten to Text")

# --- sidebar: API key choice ---
st.sidebar.header("API Key")
choice = st.sidebar.radio("Use:", ["Default API (5/day)", "My own API (unlimited)"])
if choice == "My own API (unlimited)":
    api_key = st.sidebar.text_input("Enter your Gemini API key", type="password")
    st.sidebar.markdown(
        "[See how to btain API](https://aistudio.google.com/welcome?utm_source=google&utm_medium=cpc&utm_campaign=FY25-global-DR-gsem-BKWS-1710442&utm_content=text-ad-none-any-DEV_c-CRE_736651364289-ADGP_Hybrid%20%7C%20BKWS%20-%20EXA%20%7C%20Txt-Gemini%20(Growth)-Gemini%20API-KWID_43700081658580172-aud-2306308323534:kwd-927524447508&utm_term=KW_gemini%20api-ST_gemini%20api&gclsrc=aw.ds&gad_source=1&gad_campaignid=22307834138&gbraid=0AAAAACn9t65iHO0_47zIUOq2eMGR6hDYk&gclid=CjwKCAjwy7HEBhBJEiwA5hQNokXZKrt1lM-AB05JkkRfrvPLnPpju0SUQLHMyVWU25H5t8vfb3IXJxoCgwMQAvD_BwE)"
    )
    use_default = False
else:
    api_key = DEFAULT_API_KEY
    use_default = True

# --- input/output selectors ---
input_type = st.selectbox("Select input type:", ["Image", "PDF", "Word"])
output_type = st.selectbox("Select output format:", ["TXT", "PDF", "DOCX"])

# Only for PDF/Word, show a cautionary note
if input_type in ("PDF", "Word"):
    st.warning("Make sure your documents contain **pictures** of handwriting, _not_ embedded text.")

ext_map = {
    "Image": ["png", "jpg", "jpeg", "bmp"],
    "PDF":   ["pdf"],
    "Word":  ["docx"],
}
upload = st.file_uploader(f"Upload your {input_type} file", type=ext_map[input_type])

if upload and st.button("▶️ Convert"):
    if use_default:
        usage = load_usage()
        if usage["count"] >= DAILY_LIMIT:
            st.error(f"You’ve hit the {DAILY_LIMIT}/day limit with the default API.")
            st.stop()

    if not api_key:
        st.error("❗ Please enter a valid API key.")
        st.stop()

    client = genai.Client(api_key=api_key)
    prompt = (
        f"Please extract all handwritten text from this {input_type} exactly as it appears, "
        "without adding or removing anything."
    )

    with st.spinner("Running OCR…"):
        try:
            raw = upload.read()
            texts = []

            if input_type == "Image":
                texts.append(ocr_with_gemini(client, raw, upload.name, prompt))

            elif input_type == "PDF":
                # Render each PDF page to a PNG with PyMuPDF (no Poppler needed)
                with fitz.open(stream=raw, filetype="pdf") as doc:
                    if doc.page_count == 0:
                        st.error("No pages found in the PDF.")
                        st.stop()

                    # 200 DPI is usually enough for handwriting & faster/cheaper
                    zoom = 200 / 72
                    mat = fitz.Matrix(zoom, zoom)
                    prog = st.progress(0.0)
                    total = doc.page_count

                    for i, page in enumerate(doc, start=1):
                        try:
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            img_bytes = pix.tobytes("png")
                            text_i = ocr_with_gemini(
                                client, img_bytes, f"page-{i}.png", prompt
                            )
                            texts.append(text_i)
                        except Exception as e:
                            # Don’t kill the whole run if one page is stubborn
                            texts.append(f"[Page {i} failed after retries: {e}]")
                        finally:
                            # small pacing helps reduce 503s
                            time.sleep(0.2)
                            prog.progress(i / total)

            else:  # Word (.docx)
                # Extract embedded images and OCR them
                z = zipfile.ZipFile(BytesIO(raw))
                imgs = [n for n in z.namelist() if n.startswith("word/media/")]
                if not imgs:
                    st.error("No embedded images found in the Word document.")
                    st.stop()
                for i, name in enumerate(imgs, start=1):
                    img_bytes = z.read(name)
                    try:
                        texts.append(ocr_with_gemini(client, img_bytes, os.path.basename(name), prompt))
                    except Exception as e:
                        texts.append(f"[Image {i} failed after retries: {e}]")
                    time.sleep(0.2)

            full_text = "\n\n--- Page Break ---\n\n".join(t.strip() for t in texts)

            if not full_text.strip():
                st.warning("⚠️ No text detected.")
            else:
                st.success("✅ Extracted text:")
                st.text_area("", full_text, height=250)

                # prepare output
                if output_type == "TXT":
                    data_out, ext, mime = full_text.encode("utf-8"), "txt", "text/plain"
                elif output_type == "PDF":
                    data_out, ext, mime = make_pdf(full_text), "pdf", "application/pdf"
                else:  # DOCX
                    data_out, ext, mime = make_docx(full_text), "docx", \
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                st.download_button(
                    f"📥 Download .{ext}",
                    data=data_out,
                    file_name=f"output.{ext}",
                    mime=mime,
                )

                if use_default:
                    usage["count"] += 1
                    save_usage(usage)
                    st.info(f"✅ Default-API usage: {usage['count']}/{DAILY_LIMIT}")

        except Exception as e:
            st.error(f"❌ Error: {e}")




# import os
# import json
# import mimetypes
# import zipfile
# from datetime import date
# from io import BytesIO
# from google import genai

# import streamlit as st
# from google import genai
# from google.genai import types
# import time, random, hashlib

# # for PDF → image
# # from pdf2image import convert_from_bytes  # pip install pdf2image pillow
# import fitz
# # for DOCX output
# from docx import Document             # pip install python-docx
# # for PDF output
# from fpdf import FPDF                 # pip install fpdf

# # ←── PAGE & THEME CONFIG ──────────────────────────────────────────────────────
# st.set_page_config(
#     page_title="Handwritten OCR",
#     page_icon="🖋️",
#     layout="centered",
#     initial_sidebar_state="expanded"
# )

# st.markdown(
#     """
#     <style>
#       :root {--primary-color:#004080;--bg-color:#f0f2f6;--sec-bg-color:#ffffff;--text-color:#333;}
#       .main, .reportview-container {background-color:var(--bg-color)!important;color:var(--text-color)!important;}
#       .block-container {padding:1.5rem 2rem!important;}
#       h1 {font-size:3rem!important;color:var(--primary-color)!important;}
#       h2 {font-size:2rem!important;color:var(--primary-color)!important;}
#       .stMarkdown, .stText {font-size:18px!important;}
#       .stTextArea>div>div>textarea {font-size:16px!important;}
#       .stButton>button {
#         background-color:var(--primary-color)!important;
#         color:#fff!important;
#         font-size:16px!important;
#         padding:0.5rem 1rem!important;
#       }
#       .sidebar .sidebar-content {
#         background-color:var(--sec-bg-color)!important;
#         color:var(--text-color)!important;
#         font-size:18px!important;
#       }
#     </style>
#     """,
#     unsafe_allow_html=True,
# )

# # ←── API & USAGE ───────────────────────────────────────────────────────────────
# # DEFAULT_API_KEY = " "
# DEFAULT_API_KEY = st.secrets["genai"]["api_key"]  
# USAGE_FILE = "usage.json"
# DAILY_LIMIT = 5

# def guess_mime(fname: str) -> str:
#     mime, _ = mimetypes.guess_type(fname)
#     return mime or "application/octet-stream"

# def ocr_with_gemini(client: genai.Client, data: bytes, fname: str, prompt: str) -> str:
#     part = types.Part.from_bytes(data=data, mime_type=guess_mime(fname))
#     resp = client.models.generate_content(
#         model="gemini-2.5-flash",
#         contents=[part, prompt],
#     )
#     return resp.text

# def load_usage():
#     if os.path.isfile(USAGE_FILE):
#         with open(USAGE_FILE, "r") as f:
#             d = json.load(f)
#     else:
#         d = {"date": "", "count": 0}
#     today = date.today().isoformat()
#     if d.get("date") != today:
#         d = {"date": today, "count": 0}
#     return d

# def save_usage(d):
#     with open(USAGE_FILE, "w") as f:
#         json.dump(d, f)

# def make_pdf(text: str) -> bytes:
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.set_font("Arial", size=12)
#     for line in text.splitlines():
#         pdf.multi_cell(0, 10, line)
#     return pdf.output(dest="S").encode("latin1")

# def make_docx(text: str) -> bytes:
#     doc = Document()
#     for line in text.splitlines():
#         doc.add_paragraph(line)
#     buf = BytesIO()
#     doc.save(buf)
#     return buf.getvalue()

# st.title("Handwritten to Text")

# # --- sidebar: API key choice ---
# st.sidebar.header("API Key")
# choice = st.sidebar.radio("Use:", ["Default API (5/day)", "My own API (unlimited)"])
# if choice == "My own API (unlimited)":
#     api_key = st.sidebar.text_input("Enter your Gemini API key", type="password")
#     # ←─ Display the “See how to btain API” link ───────────────
#     st.sidebar.markdown(
#         "[See how to btain API](https://aistudio.google.com/welcome?utm_source=google&utm_medium=cpc&utm_campaign=FY25-global-DR-gsem-BKWS-1710442&utm_content=text-ad-none-any-DEV_c-CRE_736651364289-ADGP_Hybrid%20%7C%20BKWS%20-%20EXA%20%7C%20Txt-Gemini%20(Growth)-Gemini%20API-KWID_43700081658580172-aud-2306308323534:kwd-927524447508&utm_term=KW_gemini%20api-ST_gemini%20api&gclsrc=aw.ds&gad_source=1&gad_campaignid=22307834138&gbraid=0AAAAACn9t65iHO0_47zIUOq2eMGR6hDYk&gclid=CjwKCAjwy7HEBhBJEiwA5hQNokXZKrt1lM-AB05JkkRfrvPLnPpju0SUQLHMyVWU25H5t8vfb3IXJxoCgwMQAvD_BwE)"
#     )
#     use_default = False
# else:
#     api_key = DEFAULT_API_KEY
#     use_default = True

# # --- input/output selectors ---
# input_type = st.selectbox("Select input type:", ["Image", "PDF", "Word"])
# output_type = st.selectbox("Select output format:", ["TXT", "PDF", "DOCX"])

# # Only for PDF/Word, show a cautionary note
# if input_type in ("PDF", "Word"):
#     st.warning("Make sure your documents contain **pictures** of handwriting, _not_ embedded text.")

# ext_map = {
#     "Image": ["png", "jpg", "jpeg", "bmp"],
#     "PDF":   ["pdf"],
#     "Word":  ["docx"],
# }
# upload = st.file_uploader(f"Upload your {input_type} file", type=ext_map[input_type])

# if upload and st.button("▶️ Convert"):
#     if use_default:
#         usage = load_usage()
#         if usage["count"] >= DAILY_LIMIT:
#             st.error(f"You’ve hit the {DAILY_LIMIT}/day limit with the default API.")
#             st.stop()

#     if not api_key:
#         st.error("❗ Please enter a valid API key.")
#         st.stop()

#     client = genai.Client(api_key=api_key)
#     prompt = (
#         f"Please extract all handwritten text from this {input_type} exactly as it appears, "
#         "without adding or removing anything."
#     )

#     with st.spinner("Running OCR…"):
#         try:
#             raw = upload.read()
#             texts = []

#             if input_type == "Image":
#                 texts.append(ocr_with_gemini(client, raw, upload.name, prompt))

#             # elif input_type == "PDF":
#             #     pages = convert_from_bytes(raw, dpi=300)
#             #     for i, page in enumerate(pages, start=1):
#             #         buf = BytesIO()
#             #         page.save(buf, format="PNG")

#             elif input_type == "PDF":
#                 # Render each PDF page to a PNG (no external binaries needed)
#                 with fitz.open(stream=raw, filetype="pdf") as doc:
#                     if doc.page_count == 0:
#                         st.error("No pages found in the PDF.")
#                         st.stop()
#                     for i, page in enumerate(doc, start=1):
#                         # 300 dpi rendering: zoom factor = dpi / 72
#                         zoom = 300 / 72
#                         mat = fitz.Matrix(zoom, zoom)
#                         pix = page.get_pixmap(matrix=mat, alpha=False)
#                         img_bytes = pix.tobytes("png")
#                         texts.append(
#                             ocr_with_gemini(client, img_bytes, f"page-{i}.png", prompt)
#                         )

            
#             else:  # Word (.docx)
#                 z = zipfile.ZipFile(BytesIO(raw))
#                 imgs = [n for n in z.namelist() if n.startswith("word/media/")]
#                 if not imgs:
#                     st.error("No embedded images found in the Word document.")
#                     st.stop()
#                 for name in imgs:
#                     img_bytes = z.read(name)
#                     texts.append(ocr_with_gemini(client, img_bytes, os.path.basename(name), prompt))

#             full_text = "\n\n--- Page Break ---\n\n".join(t.strip() for t in texts)

#             if not full_text.strip():
#                 st.warning("⚠️ No text detected.")
#             else:
#                 st.success("✅ Extracted text:")
#                 st.text_area("", full_text, height=250)

#                 # prepare output
#                 if output_type == "TXT":
#                     data_out, ext, mime = full_text.encode("utf-8"), "txt", "text/plain"
#                 elif output_type == "PDF":
#                     data_out, ext, mime = make_pdf(full_text), "pdf", "application/pdf"
#                 else:  # DOCX
#                     data_out, ext, mime = make_docx(full_text), "docx", \
#                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#                 st.download_button(
#                     f"📥 Download .{ext}",
#                     data=data_out,
#                     file_name=f"output.{ext}",
#                     mime=mime,
#                 )

#                 if use_default:
#                     usage["count"] += 1
#                     save_usage(usage)
#                     st.info(f"✅ Default-API usage: {usage['count']}/{DAILY_LIMIT}")

#         except Exception as e:
#             st.error(f"❌ Error: {e}")


